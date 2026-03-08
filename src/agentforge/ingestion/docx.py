"""Word document (.docx) job description ingestion."""

from __future__ import annotations

from pathlib import Path

from agentforge.ingestion.text import ingest_text
from agentforge.models.job_description import JDSource, JobDescription


def ingest_docx(path: str | Path, company: str | None = None) -> JobDescription:
    """Extract text from a .docx file and ingest as a job description."""
    from docx import Document

    file_path = Path(path)
    if not file_path.exists():
        raise FileNotFoundError(f"DOCX file not found: {path}")

    try:
        doc = Document(str(file_path))
    except Exception as e:
        raise ValueError(f"Failed to open DOCX file: {path}. File may be corrupted.") from e

    # Extract paragraphs
    paragraphs: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        # Preserve heading structure with markdown-style markers
        if para.style and para.style.name.startswith("Heading"):
            level = para.style.name.replace("Heading ", "").replace("Heading", "1")
            try:
                hashes = "#" * int(level)
            except ValueError:
                hashes = "#"
            paragraphs.append(f"{hashes} {text}")
        else:
            paragraphs.append(text)

    # Extract text from tables
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))

    full_text = "\n\n".join(paragraphs)
    if not full_text.strip():
        raise ValueError(f"No text content extracted from DOCX: {path}")

    jd = ingest_text(full_text, company=company)
    jd.source = JDSource.FILE
    jd.metadata["file_path"] = str(file_path.resolve())
    jd.metadata["format"] = "docx"
    return jd
