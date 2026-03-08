"""Tests for DOCX ingestion."""

from __future__ import annotations

from pathlib import Path

import pytest

from agentforge.ingestion.docx import ingest_docx
from agentforge.models.job_description import JDSource


def _create_test_docx(path: Path, paragraphs: list[str], heading: str | None = None) -> Path:
    """Create a simple test .docx file."""
    from docx import Document

    doc = Document()
    if heading:
        doc.add_heading(heading, level=1)
    for para in paragraphs:
        doc.add_paragraph(para)
    doc.save(str(path))
    return path


class TestDocxIngestion:
    def test_basic_extraction(self, tmp_path):
        docx_path = tmp_path / "test.docx"
        _create_test_docx(
            docx_path,
            ["We are looking for a software engineer.", "Must have 3+ years experience."],
            heading="Software Engineer",
        )

        jd = ingest_docx(docx_path)
        assert "software engineer" in jd.raw_text.lower()
        assert jd.source == JDSource.FILE
        assert jd.metadata["format"] == "docx"

    def test_heading_preserved(self, tmp_path):
        docx_path = tmp_path / "test.docx"
        _create_test_docx(
            docx_path,
            ["Design data pipelines."],
            heading="Data Engineer",
        )

        jd = ingest_docx(docx_path)
        # Heading should be converted to markdown-style
        assert "# Data Engineer" in jd.raw_text

    def test_file_not_found(self):
        with pytest.raises(FileNotFoundError):
            ingest_docx("/nonexistent/file.docx")

    def test_empty_document(self, tmp_path):
        docx_path = tmp_path / "empty.docx"
        from docx import Document
        doc = Document()
        doc.save(str(docx_path))

        with pytest.raises(ValueError, match="No text content"):
            ingest_docx(docx_path)

    def test_with_company(self, tmp_path):
        docx_path = tmp_path / "test.docx"
        _create_test_docx(docx_path, ["Looking for a manager."])

        jd = ingest_docx(docx_path, company="Acme Corp")
        assert jd.company == "Acme Corp"

    def test_file_path_in_metadata(self, tmp_path):
        docx_path = tmp_path / "test.docx"
        _create_test_docx(docx_path, ["Test content here."])

        jd = ingest_docx(docx_path)
        assert "file_path" in jd.metadata
        assert str(docx_path.resolve()) in jd.metadata["file_path"]

    def test_table_extraction(self, tmp_path):
        """Tables in docx should be extracted as pipe-separated rows."""
        from docx import Document

        docx_path = tmp_path / "table.docx"
        doc = Document()
        doc.add_paragraph("Job Requirements")
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Skill"
        table.cell(0, 1).text = "Level"
        table.cell(1, 0).text = "Python"
        table.cell(1, 1).text = "Expert"
        doc.save(str(docx_path))

        jd = ingest_docx(docx_path)
        assert "Python" in jd.raw_text
        assert "|" in jd.raw_text

    def test_corrupted_file(self, tmp_path):
        """Corrupted docx should raise ValueError."""
        bad_path = tmp_path / "bad.docx"
        bad_path.write_text("not a valid docx file")

        with pytest.raises(ValueError, match="Failed to open DOCX"):
            ingest_docx(bad_path)

    def test_heading_level_2(self, tmp_path):
        """Level 2 headings should produce ## markers."""
        from docx import Document

        docx_path = tmp_path / "heading.docx"
        doc = Document()
        doc.add_heading("Subsection Title", level=2)
        doc.add_paragraph("Content under subsection")
        doc.save(str(docx_path))

        jd = ingest_docx(docx_path)
        assert "## Subsection Title" in jd.raw_text
